"""Read text/data from Word (.docx) and Excel (.xlsx/.xlsm) files."""

SCHEMA = {
    "name": "office_reader",
    "description": "Read text from a .docx or cell data from a .xlsx/.xlsm file.",
    "parameters": {"path": {"type": "string"}},
    "required": ["path"],
}


def handle(args: dict) -> dict:
    path = args.get("path", "")
    if not path:
        return {"error": "missing_path"}
    low = path.lower()
    try:
        if low.endswith(".docx"):
            import docx
            doc = docx.Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            return {"type": "docx", "paragraphs": len(doc.paragraphs),
                    "text": text[:8000], "truncated": len(text) > 8000}
        if low.endswith((".xlsx", ".xlsm")):
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            sheets = {}
            for ws in wb.worksheets[:5]:
                rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= 50:
                        break
                    rows.append([("" if c is None else c) for c in row])
                sheets[ws.title] = {"dimensions": ws.dimensions, "rows": rows}
            wb.close()
            return {"type": "xlsx", "sheet_count": len(sheets), "sheets": sheets}
        return {"error": "unsupported_format", "path": path}
    except Exception as e:
        return {"error": str(e), "path": path}
